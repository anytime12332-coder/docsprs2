"""Multi-format document parser."""

import io
from typing import Any

from loguru import logger


class DocumentParser:
    """Parses various document formats and extracts text content."""

    async def parse(self, content: bytes, mime_type: str, filename: str) -> dict[str, Any]:
        """Parse document based on mime type."""
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

        if mime_type == "application/pdf" or ext == "pdf":
            return await self._parse_pdf(content)
        elif mime_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ) or ext == "docx":
            return await self._parse_docx(content)
        elif mime_type in (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ) or ext == "xlsx":
            return await self._parse_xlsx(content)
        elif ext == "csv" or mime_type == "text/csv":
            return await self._parse_csv(content)
        elif ext == "html" or mime_type == "text/html":
            return await self._parse_html(content)
        elif mime_type.startswith("image/") or ext in (
            "png", "jpg", "jpeg", "tiff", "bmp", "webp"
        ):
            return {"type": "image", "text": "", "pages": [], "requires_ocr": True}
        else:
            raise ValueError(f"Unsupported format: {mime_type} ({ext})")

    async def _parse_pdf(self, content: bytes) -> dict[str, Any]:
        from app.processing.pdf_processor import pdf_processor

        pages = await pdf_processor.extract_pages(content)
        page_count = len(pages)
        full_text = "\n\n".join(p["text"] for p in pages if p["text"])
        metadata = await pdf_processor.extract_metadata(content)

        # Check if OCR is needed (pages with no text)
        needs_ocr = any(not p["text"].strip() for p in pages)

        all_tables = []
        for p in pages:
            for t in p.get("tables", []):
                t["page_number"] = p["page_number"]
                all_tables.append(t)

        return {
            "type": "pdf",
            "text": full_text,
            "pages": pages,
            "page_count": page_count,
            "tables": all_tables,
            "metadata": metadata,
            "requires_ocr": needs_ocr,
        }

    async def _parse_docx(self, content: bytes) -> dict[str, Any]:
        from docx import Document

        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)

        tables = []
        for table in doc.tables:
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            rows = []
            for row in table.rows[1:]:
                rows.append([cell.text.strip() for cell in row.cells])
            tables.append({"headers": headers, "rows": rows})

        return {
            "type": "docx",
            "text": full_text,
            "pages": [{"page_number": 1, "text": full_text}],
            "page_count": 1,
            "tables": tables,
            "requires_ocr": False,
        }

    async def _parse_xlsx(self, content: bytes) -> dict[str, Any]:
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        tables = []
        all_text = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows_data = []
            headers = []

            for i, row in enumerate(ws.iter_rows(values_only=True)):
                row_values = [str(cell) if cell is not None else "" for cell in row]
                if i == 0:
                    headers = row_values
                else:
                    rows_data.append(row_values)
                all_text.extend(row_values)

            tables.append(
                {
                    "table_name": sheet_name,
                    "headers": headers,
                    "rows": rows_data,
                }
            )

        wb.close()

        return {
            "type": "xlsx",
            "text": " ".join(filter(None, all_text)),
            "pages": [],
            "page_count": len(wb.sheetnames),
            "tables": tables,
            "requires_ocr": False,
        }

    async def _parse_csv(self, content: bytes) -> dict[str, Any]:
        import csv

        text = content.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(text))
        rows = list(reader)

        headers = rows[0] if rows else []
        data_rows = rows[1:] if len(rows) > 1 else []

        return {
            "type": "csv",
            "text": text,
            "pages": [],
            "page_count": 1,
            "tables": [{"headers": headers, "rows": data_rows}],
            "requires_ocr": False,
        }

    async def _parse_html(self, content: bytes) -> dict[str, Any]:
        from html.parser import HTMLParser

        class TextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self.text_parts = []

            def handle_data(self, data):
                text = data.strip()
                if text:
                    self.text_parts.append(text)

        html_text = content.decode("utf-8", errors="replace")
        extractor = TextExtractor()
        extractor.feed(html_text)
        text = " ".join(extractor.text_parts)

        return {
            "type": "html",
            "text": text,
            "pages": [{"page_number": 1, "text": text}],
            "page_count": 1,
            "tables": [],
            "requires_ocr": False,
        }


document_parser = DocumentParser()
